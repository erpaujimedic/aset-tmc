from fastapi import APIRouter, HTTPException, Query, UploadFile, File
from fastapi.responses import StreamingResponse
from app.database import supabase
from pydantic import BaseModel
from typing import Optional, List, Any
from fastapi_cache.decorator import cache
import io
import openpyxl
import uuid

router = APIRouter(prefix="/assets", tags=["Assets"])

class AssetCreate(BaseModel):
    id: str
    name: str
    category: str
    branch: str
    department: str
    room: str
    condition: str
    photo_url: str
    brand: Optional[str] = None
    serial_number: Optional[str] = None
    user_name: Optional[str] = None
    user_phone: Optional[str] = None
    pr_number: Optional[str] = None
    placement_location: Optional[str] = None
    rack_number: Optional[str] = None
    calibration_doc_url: Optional[str] = None
    assignee: Optional[str] = None
    status: str = "Active"
    lat: Optional[float] = None
    lng: Optional[float] = None
    location_name: Optional[str] = None
    is_labeled: bool = False

class AssetUpdate(BaseModel):
    name: Optional[str] = None
    category: Optional[str] = None
    branch: Optional[str] = None
    department: Optional[str] = None
    room: Optional[str] = None
    condition: Optional[str] = None
    photo_url: Optional[str] = None
    brand: Optional[str] = None
    serial_number: Optional[str] = None
    user_name: Optional[str] = None
    user_phone: Optional[str] = None
    pr_number: Optional[str] = None
    placement_location: Optional[str] = None
    rack_number: Optional[str] = None
    calibration_doc_url: Optional[str] = None
    assignee: Optional[str] = None
    status: Optional[str] = None
    lat: Optional[float] = None
    lng: Optional[float] = None
    location_name: Optional[str] = None
    is_labeled: Optional[bool] = None

from fastapi_cache import FastAPICache
import base64
from docxtpl import DocxTemplate, InlineImage
from docx.shared import Mm
import docx
from copy import deepcopy
import qrcode

class AssetBARow(BaseModel):
    asset_id: str
    no_pr: str
    branch: str
    nama_asset: str
    serial_number: str
    user: str
    department: str
    is_labeled: bool = False

class BAPayload(BaseModel):
    tanggal: str
    assets: List[AssetBARow]

@router.post("/export-ba")
async def export_ba(payload: BAPayload):
    try:
        doc = docx.Document("templates/ba_barcodeassets.docx")
    
    target_row = None
    target_table = None
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if '{{nama_asset}}' in cell.text:
                    target_row = row
                    target_table = table
                    break
            if target_row:
                break
        if target_row:
            break
            
    if target_row and len(payload.assets) > 0:
        for cell in target_row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.text = run.text.replace("}}", "_0}}")
                    
        prev_row_tr = target_row._tr
        for i in range(1, len(payload.assets)):
            new_row_tr = deepcopy(target_row._tr)
            prev_row_tr.addnext(new_row_tr)
            new_row = docx.table._Row(new_row_tr, target_table)
            for cell_idx, cell in enumerate(new_row.cells):
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.text = run.text.replace("_0}}", "_" + str(i) + "}}")
                        if cell_idx == 0 and "1" in run.text:
                            run.text = run.text.replace("1", str(i + 1))
            prev_row_tr = new_row_tr
            
    temp_buf = io.BytesIO()
    doc.save(temp_buf)
    temp_buf.seek(0)
    
    doc_tpl = DocxTemplate(temp_buf)
    
    unlabeled = [a.nama_asset for a in payload.assets if not a.is_labeled]
    if unlabeled:
        ket_barcode = f"{', '.join(unlabeled)} belum terlabeli."
    else:
        ket_barcode = "-"
        
    context = {
        "tanggal": payload.tanggal,
        "ket_barcode": ket_barcode
    }
    
    from PIL import Image, ImageDraw, ImageFont
    import tempfile
    import os
    
    frontend_logo_path = "eam-frontend/public/logo.png"
    if not os.path.exists(frontend_logo_path):
        frontend_logo_path = "../eam-frontend/public/logo.png"
        
    for i, asset in enumerate(payload.assets):
        context[f"no_pr_{i}"] = asset.no_pr if asset.no_pr else "-"
        context[f"branch_{i}"] = asset.branch if asset.branch else "-"
        context[f"nama_asset_{i}"] = asset.nama_asset if asset.nama_asset else "-"
        context[f"serial_number_{i}"] = asset.serial_number if asset.serial_number and asset.serial_number != "NN" else "-"
        context[f"department_{i}"] = asset.department if asset.department else "-"
        
        try:
            scale = 4
            qr = qrcode.QRCode(version=1, box_size=10 * scale, border=1)
            qr.add_data(f"http://localhost:5173/public-asset/{asset.asset_id}")
            qr.make(fit=True)
            qr_img = qr.make_image(fill_color="black", back_color="white").convert('RGB')
            
            qr_size = 180 * scale
            qr_img = qr_img.resize((qr_size, qr_size))
            
            canvas_w = 200 * scale
            canvas_h = 240 * scale
            canvas = Image.new('RGB', (canvas_w, canvas_h), 'white')
            canvas.paste(qr_img, (10 * scale, 5 * scale))
            
            draw = ImageDraw.Draw(canvas)
            try:
                font_bold_small = ImageFont.truetype('arialbd.ttf', 16 * scale)
                font_bold_tiny = ImageFont.truetype('arialbd.ttf', 14 * scale)
            except:
                font_bold_small = ImageFont.load_default()
                font_bold_tiny = ImageFont.load_default()
            
            if os.path.exists(frontend_logo_path):
                logo = Image.open(frontend_logo_path).convert('RGBA')
                lw, lh = logo.size
                new_h = 16 * scale
                new_w = int(lw * (new_h / lh))
                logo = logo.resize((new_w, new_h))
                logo_exists = True
            else:
                logo_exists = False
                new_w = 0
                
            name_text = asset.nama_asset.upper()
            if len(name_text) > 15:
                name_text = name_text[:13] + ".."
                
            name_bbox = font_bold_tiny.getbbox(name_text)
            name_w = name_bbox[2] - name_bbox[0]
            
            total_w = (new_w + (4 * scale) + name_w) if logo_exists else name_w
            start_x = (canvas_w - total_w) / 2
            
            text_y1 = 195 * scale
            if logo_exists:
                canvas.paste(logo, (int(start_x), text_y1), logo)
                text_x = start_x + new_w + (4 * scale)
            else:
                text_x = start_x
                
            draw.text((text_x, text_y1), name_text, fill='black', font=font_bold_tiny)
            
            id_bbox = font_bold_small.getbbox(asset.asset_id)
            id_w = id_bbox[2] - id_bbox[0]
            id_x = (canvas_w - id_w) / 2
            text_y2 = 215 * scale
            draw.text((id_x, text_y2), asset.asset_id, fill='black', font=font_bold_small)
            
            img_stream = io.BytesIO()
            canvas.save(img_stream, format='PNG')
            img_stream.seek(0)
            context[f"qr_code_{i}"] = InlineImage(doc_tpl, img_stream, width=Mm(18))
        except Exception as e:
            print("QR Error:", e)
            context[f"qr_code_{i}"] = ""
            
    doc_tpl.render(context)
    
    with tempfile.TemporaryDirectory() as tmpdir:
        temp_docx = os.path.join(tmpdir, "temp.docx")
        temp_pdf = os.path.join(tmpdir, "temp.pdf")
        doc_tpl.save(temp_docx)
        
        try:
            import sys
            import subprocess
            import os
            python_exe = os.path.join(os.getcwd(), ".venv", "Scripts", "python.exe")
            if not os.path.exists(python_exe):
                python_exe = sys.executable
            
            script = f"from docx2pdf import convert; convert(r'{temp_docx}', r'{temp_pdf}')"
            result = subprocess.run([python_exe, "-c", script], capture_output=True, text=True)
            if result.returncode != 0:
                print("PDF Subprocess Error:", result.stderr)
                raise Exception("Subprocess convert failed")
            
            with open(temp_pdf, "rb") as f:
                pdf_data = f.read()
            out_buf = io.BytesIO(pdf_data)
            media_type = "application/pdf"
            filename = "Berita_Acara_Asset.pdf"
        except Exception as e:
            print("PDF Convert Error:", e)
            out_buf = io.BytesIO()
            doc_tpl.save(out_buf)
            out_buf.seek(0)
            media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            filename = "Berita_Acara_Asset.docx"
    
        return StreamingResponse(
            out_buf,
            media_type=media_type,
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
    except Exception as e:
        import traceback
        from fastapi.responses import PlainTextResponse
        return PlainTextResponse(content=traceback.format_exc(), status_code=500)

class ExportData(BaseModel):
    headers: List[str]
    rows: List[List[Any]]

@router.post("/export-excel")
async def export_excel(data: ExportData):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Assets Export"
    
    ws.append(data.headers)
    for cell in ws[1]:
        cell.font = openpyxl.styles.Font(bold=True, color="FFFFFF")
        cell.fill = openpyxl.styles.PatternFill(start_color="286086", end_color="286086", fill_type="solid")
    
    for row in data.rows:
        ws.append(row)
        
    for col in ws.columns:
        max_length = 0
        column = col[0].column_letter
        for cell in col:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            except:
                pass
        ws.column_dimensions[column].width = max_length + 2
        
    ws.freeze_panes = "A2"
    
    out = io.BytesIO()
    wb.save(out)
    out.seek(0)
    
    return StreamingResponse(
        out,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": "attachment; filename=assets_inventory.xlsx"}
    )


@router.get("")
@cache(expire=3600, namespace="assets")
async def get_assets(branch: Optional[str] = None, status: Optional[str] = None):
    if not supabase:
        raise HTTPException(status_code=500, detail="Database connection error")
    try:
        all_data = []
        page_size = 1000
        start = 0
        
        while True:
            query = supabase.table("assets").select("*")
            if branch and branch not in ("All Branches", "ALL", "ALL Branches"):
                query = query.eq("branch", branch)
            if status and status != "All Status":
                query = query.eq("status", status)
                
            res = query.range(start, start + page_size - 1).execute()
            
            if not res.data:
                break
                
            all_data.extend(res.data)
            
            if len(res.data) < page_size:
                break
                
            start += page_size
            
        return {"data": all_data}
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))

@router.post("")
async def create_asset(asset: AssetCreate):
    if not supabase:
        raise HTTPException(status_code=500, detail="Database connection error")
    try:
        res = supabase.table("assets").insert(asset.dict()).execute()
        await FastAPICache.clear(namespace="assets")
        return {"message": "Asset created successfully", "data": res.data}
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))

@router.put("/{asset_id}")
async def update_asset(asset_id: str, asset: AssetUpdate):
    if not supabase:
        raise HTTPException(status_code=500, detail="Database connection error")
    try:
        update_data = {k: v for k, v in asset.dict().items() if v is not None}
        if not update_data:
            return {"message": "No data to update"}
        res = supabase.table("assets").update(update_data).eq("id", asset_id).execute()
        await FastAPICache.clear(namespace="assets")
        return {"message": "Asset updated successfully", "data": res.data}
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))

@router.delete("/{asset_id}")
async def delete_asset(asset_id: str):
    if not supabase:
        raise HTTPException(status_code=500, detail="Database connection error")
    try:
        res = supabase.table("assets").delete().eq("id", asset_id).execute()
        await FastAPICache.clear(namespace="assets")
        return {"message": "Asset deleted successfully"}
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))

@router.get("/import-template")
def get_import_template():
    wb = openpyxl.Workbook()
    
    # Sheet 1: Import Bulk Data
    ws1 = wb.active
    ws1.title = "Import Bulk Data"
    headers = [
        "Kode Asset (Opsional)", "Nama Asset (Wajib)", "Merk Barang", "Serial Number", "Nama User", "No. HP User",
        "Lokasi", "Department", "Kategori Asset", "No PR", "Lokasi Penempatan",
        "Nomor Rak", "Ruangan", "Kondisi", "Lampiran Foto Asset", "Lampiran Kalibrasi"
    ]
    ws1.append(headers)
    
    # Optional styling for headers
    for cell in ws1[1]:
        cell.font = openpyxl.styles.Font(bold=True)
    
    # Sheet 2: Information
    ws2 = wb.create_sheet(title="Information")
    info_data = [
        ["Informasi Pengisian Bulk Import"],
        [],
        ["KODE ASSET", "Jika dikosongkan, sistem otomatis membuat ID baru. Jika Anda memiliki kode aset sendiri, ketikkan di kolom pertama."],
        [],
        ["KATEGORI ASSET", "Kode", "Deskripsi"],
        ["", "FFF", "Furniture"],
        ["", "ELK", "Elektronik Non Alat Kesehatan"],
        ["", "ALK", "Elektronik Alat Kesehatan"],
        ["", "VH2", "Kendaraan Roda 2"],
        ["", "VH4", "Kendaraan Roda 4"],
        ["", "HRW", "Hardware"],
        ["", "LGL", "Surat Berharga"],
        ["", "PRK", "Perkakas"],
        [],
        ["KONDISI ASSET"],
        ["", "BAGUS & DIGUNAKAN"],
        ["", "BAGUS & TIDAK DIGUNAKAN"],
        ["", "RUSAK & PERLU PERGANTIAN"],
        ["", "RUSAK & PERLU DIMUSNAHKAN"],
        [],
        ["LOKASI PENEMPATAN & NOMOR RAK"],
        ["Khusus untuk Head Office, wajib mengisi Lokasi Penempatan (cth: Lantai 1, Lantai 2, Gudang)."],
        ["Jika Lokasi Penempatan adalah 'Gudang' atau 'Warehouse', wajib mengisi Nomor Rak."]
    ]
    for row in info_data:
        ws2.append(row)
        
    out = io.BytesIO()
    wb.save(out)
    out.seek(0)
    
    return StreamingResponse(
        out,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": "attachment; filename=Asset_Import_Template.xlsx"}
    )

@router.post("/import")
async def import_assets(file: UploadFile = File(...)):
    if not supabase:
        raise HTTPException(status_code=500, detail="Database connection error")
    
    if not file.filename.endswith('.xlsx'):
        raise HTTPException(status_code=400, detail="Only .xlsx files are supported")
        
    try:
        contents = await file.read()
        wb = openpyxl.load_workbook(io.BytesIO(contents), data_only=True)
        ws = wb["Import Bulk Data"]
        
        # Read headers
        rows = list(ws.iter_rows(values_only=True))
        if len(rows) < 2:
            raise HTTPException(status_code=400, detail="No data found in sheet")
            
        assets_to_insert = []
        import datetime
        current_year = datetime.datetime.now().year
        
        # Fetch master branches to normalize casing
        branches_res = supabase.table("branches").select("name").execute()
        valid_branches = {b["name"].lower(): b["name"] for b in branches_res.data} if branches_res.data else {}
        
        for idx, row in enumerate(rows[1:], start=2):
            # Skip empty rows
            if not row[1]:
                continue
                
            provided_id_raw = str(row[0]).strip() if row[0] else ""
            if provided_id_raw.lower() in ('none', '#n/a', 'n/a', 'null', '-') or provided_id_raw.startswith('#'):
                provided_id = None
            else:
                provided_id = provided_id_raw if provided_id_raw else None
            name = str(row[1]).strip() if row[1] else ""
            brand = str(row[2]).strip() if len(row) > 2 and row[2] else None
            serial_number = str(row[3]).strip() if len(row) > 3 and row[3] else None
            user_name = str(row[4]).strip() if len(row) > 4 and row[4] else "Unassigned"
            user_phone = str(row[5]).strip() if len(row) > 5 and row[5] else None
            branch_raw = str(row[6]).strip() if len(row) > 6 and row[6] else "Unknown Branch"
            branch = valid_branches.get(branch_raw.lower(), branch_raw)
            department = str(row[7]).strip().upper() if len(row) > 7 and row[7] else "-"
            category_code = str(row[8]).strip() if len(row) > 8 and row[8] else "UNC"
            pr_number = str(row[9]).strip() if len(row) > 9 and row[9] else None
            placement_location = str(row[10]).strip() if len(row) > 10 and row[10] else None
            rack_number = str(row[11]).strip() if len(row) > 11 and row[11] else None
            room = str(row[12]).strip() if len(row) > 12 and row[12] else "-"
            condition = str(row[13]).strip() if len(row) > 13 and row[13] else "BAGUS & DIGUNAKAN"
            photo_url = str(row[14]).strip() if len(row) > 14 and row[14] else ""
            calibration_doc_url = str(row[15]).strip() if len(row) > 15 and row[15] else None
            
            if provided_id:
                asset_id = provided_id
            else:
                asset_id = f"AST-{current_year}-{category_code}-{str(uuid.uuid4()).split('-')[0].upper()}"
            
            qr = qrcode.QRCode(version=1, box_size=10, border=5)
            qr.add_data(asset_id)
            qr.make(fit=True)
            img = qr.make_image(fill_color="black", back_color="white")
            buffer = io.BytesIO()
            img.save(buffer, format="PNG")
            qr_base64 = base64.b64encode(buffer.getvalue()).decode()
            
            asset_data = {
                "id": asset_id,
                "name": name,
                "category": category_code,
                "branch": branch,
                "department": department,
                "room": room,
                "condition": condition,
                "photo_url": photo_url,
                "brand": brand,
                "serial_number": serial_number,
                "assignee": user_name,
                "user_phone": user_phone,
                "pr_number": pr_number,
                "placement_location": placement_location,
                "rack_number": rack_number,
                "calibration_doc_url": calibration_doc_url,
                "qr_code": qr_base64,
                "is_labeled": False,
                "status": "Active"
            }
            assets_to_insert.append(asset_data)
            
        if assets_to_insert:
            res = supabase.table("assets").insert(assets_to_insert).execute()
            await FastAPICache.clear(namespace="assets")
            return {"message": f"Successfully imported {len(assets_to_insert)} assets", "data": res.data}
        else:
            return {"message": "No valid data to import"}
            
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to process file: {str(e)}")
