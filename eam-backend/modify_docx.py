import docx

def main():
    doc = docx.Document('templates/form_permintaan.docx')
    for table in doc.tables:
        target_row_idx = -1
        for i, row in enumerate(table.rows):
            # Check for Peminjaman row
            row_text = ''.join(c.text for c in row.cells)
            if 'Peminjaman' in row_text:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if '{{cabang_asal}}' in p.text:
                            p.text = p.text.replace('{{cabang_asal}}', '{{cabang_asal_peminjaman}}')
            # Check for Pemindahan row
            if 'Pemindahan' in row_text:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if '{{cabang_asal}}' in p.text:
                            p.text = p.text.replace('{{cabang_asal}}', '{{cabang_asal_pemindahan}}')
            
            # Find the data table row
            for cell in row.cells:
                if '{{nama_barang}}' in cell.text:
                    target_row_idx = i
                    break
                    
        if target_row_idx != -1:
            row = table.rows[target_row_idx]
            
            # Add a row for the {%tr for item in assets %}
            row_for = table.add_row()
            row._tr.addprevious(row_for._tr)
            row_for.cells[0].text = '{%tr for item in assets %}'
            
            # Add a row for the {%tr endfor %}
            row_endfor = table.add_row()
            row._tr.addnext(row_endfor._tr)
            row_endfor.cells[0].text = '{%tr endfor %}'
            
            # Replace '1' with '{{loop.index}}' and other variables with 'item.' prefix
            replacements = {
                '1': '{{loop.index}}',
                '{{nama_barang}}': '{{item.nama_barang}}',
                '{{code}}': '{{item.code}}',
                '{{merk}}': '{{item.merk}}',
                '{{sn}}': '{{item.sn}}',
                '{{qty}}': '{{item.qty}}',
                '{{harga}}': '{{item.harga}}',
                '{{kondisi}}': '{{item.kondisi}}',
                '{{tgl_perbaikan_sebelumnya}}': '{{item.tgl_perbaikan_sebelumnya}}',
                '{{due_date_kalibrasi}}': '{{item.due_date_kalibrasi}}',
                '{{kronologis}}': '{{item.kronologis}}'
            }
            
            for cell in row.cells:
                for p in cell.paragraphs:
                    for old_val, new_val in replacements.items():
                        if old_val in p.text:
                            p.text = p.text.replace(old_val, new_val)
                        
    doc.save('templates/form_permintaan_template.docx')
    print("Done")

if __name__ == '__main__':
    main()
